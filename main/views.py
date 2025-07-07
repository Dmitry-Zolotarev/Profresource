<<<<<<< HEAD
=======
from sqlite3 import IntegrityError

>>>>>>> 2889caa (07.07.2025)
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
<<<<<<< HEAD
=======
from django.views.decorators.http import require_POST
>>>>>>> 2889caa (07.07.2025)
from django.db.models import F, Q
import json, re
import zipfile
from io import BytesIO
from datetime import date
from datetime import datetime
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font
from openpyxl.utils import get_column_letter
import openpyxl
import hashlib

from .models import (
    Слушатели, Пол, Группы, Курсы, Типы_курсов,
    Материалы_курсов, Типы_материалов,
    Организации, Человек_группа, Человек_организация,
<<<<<<< HEAD
    Статусы, Месяцы, Страны, Приказы, Протоколы
=======
    Статусы, Месяцы, Страны, Приказы, Протоколы,
    Регионы, Районы, Населенные_пункты,
    Почтовые_индексы, Названия_улиц, Улицы, Адреса_слушателей, Адреса, Адреса_организаций
>>>>>>> 2889caa (07.07.2025)
)
def слушатели_view(request):
    if request.method == 'POST':
        error_found = False
        # Проверки только для граждан РФ (id = 1)
        if request.POST.get('Гражданство') == "1":
            if not re.fullmatch(r'\d{4}', request.POST.get('Серия_паспорта') or ''):
                messages.warning(request, "Серия паспорта РФ должна состоять из 4 цифр.")
                error_found = True
            if not re.fullmatch(r'\d{6}', request.POST.get('Номер_паспорта') or ''):
                messages.warning(request, "Номер паспорта РФ должен состоять из 6 цифр.")
                error_found = True
        if Слушатели.objects.filter(ИНН=request.POST.get('ИНН')).exists():
            messages.warning(request, "Человек с этим ИНН уже есть в базе данных!")
            error_found = True
        if not error_found:
            Слушатели(
                Фамилия=request.POST.get('Фамилия'),
                Имя=request.POST.get('Имя'),
                Отчество=request.POST.get('Отчество'),
                Дата_рождения=request.POST.get('Дата_рождения'),
                Пол_id=request.POST.get('Пол'),
                Гражданство_id=request.POST.get('Гражданство'),
                Серия_паспорта=request.POST.get('Серия_паспорта'),
                Номер_паспорта=request.POST.get('Номер_паспорта'),
                Номер_СНИЛС=request.POST.get('Номер_СНИЛС'),
                ИНН=request.POST.get('ИНН'),
                Телефон=request.POST.get('Телефон'),
                Email=request.POST.get('Email'),
            ).save()
            return redirect('student_list')
    return render(request, 'main/ListenerList.html', {
        'слушатели': Слушатели.objects.all(),
        'страны': Страны.objects.all(),
        'полы': Пол.objects.all(),
    })
# --- Группы ---
def группы_view(request):
    if request.method == 'POST':
        if 'submit_группа' in request.POST:
<<<<<<< HEAD
            if Группы.objects.filter(
                    Курс_id=request.POST.get('Курс'),
                    Дата_начала_курса=request.POST.get('Дата_начала_курса'),
                    Дата_окончания_курса=request.POST.get('Дата_окончания_курса')
            ).exists():
                messages.warning(request, "Такой курс уже есть в базе данных!")
            else:
                Группы.objects.create(
                    Курс_id=request.POST.get('Курс'),
                    Дата_начала_курса=request.POST.get('Дата_начала_курса'),
                    Дата_окончания_курса=request.POST.get('Дата_окончания_курса')
                )
=======
            if Группы.objects.filter(Курс_id=request.POST.get('Курс')).exists():
                messages.warning(request, "Такой курс уже есть в базе данных!")
            else:
                Группы.objects.create(Курс_id=request.POST.get('Курс'))
>>>>>>> 2889caa (07.07.2025)
                return redirect('group_list')
        elif 'submit_привязка' in request.POST:
            if Человек_группа.objects.filter(
                Слушатель_id=request.POST.get('Слушатель'),
                Группа_id=request.POST.get('Группа'),
            ).exists():
                messages.warning(request, "Эта привязка уже есть в базе данных!")
            else:
                Человек_группа.objects.create(
                    Слушатель_id=request.POST.get('Слушатель'),
                    Группа_id=request.POST.get('Группа'),
                    Статус_id=1,
                )
                return redirect('group_list')
    return render(request, 'main/GroupList.html', {
        'курсы': Курсы.objects.all(),
        'слушатели': Слушатели.objects.all(),
        'группы': Группы.objects.all(),
        'Человек_группа': Человек_группа.objects.all(),
        'статусы': Статусы.objects.all(),
    })
# --- Курсы ---
def курсы_view(request):
    if request.method == 'POST':
        if Курсы.objects.filter(
                Тип_id=request.POST.get('Тип'),
                Название=request.POST.get('Название'),
        ).exists():
            messages.warning(request, "Такой курс уже есть в базе данных!")
        else:
            Курсы.objects.create(
                Тип_id=request.POST.get('Тип'),
                Название=request.POST.get('Название'),
                Объём_часов=request.POST.get('Объём_часов'),
            )
            return redirect('course_list')
    return render(request, 'main/CourseList.html', {
        'типы_курсов': Типы_курсов.objects.all(),
        'курсы': Курсы.objects.all(),
    })
def страны_view(request):
    if request.method == 'POST':
        if Страны.objects.filter(
                Краткое_название=request.POST.get('Краткое_название'),
                Телефонный_код=request.POST.get('Телефонный_код'),
        ).exists():
            messages.warning(request, "Такой курс уже есть в базе данных!")
        else:
            Страны.objects.create(
                Название=request.POST.get('Название'),
                Краткое_название=request.POST.get('Краткое_название'),
                Телефонный_код=request.POST.get('Телефонный_код'),
            )
            return redirect('country_list')
    return render(request, 'main/CountryList.html', {
        'страны': Страны.objects.all(),
    })

# --- Материалы ---
def материалы_view(request):
    if request.method == 'POST':
        if Материалы_курсов.objects.filter(
                Название=request.POST.get('Название'),
                Ссылка_на_материал=request.POST.get('Ссылка'),
                Курс_id=request.POST.get('Курс'),
                Тип_id=request.POST.get('Тип'),
            ).exists():
            messages.warning(request, "Такой материал уже есть в базе данных!")
        else:
            Материалы_курсов.objects.create(
                Название=request.POST.get('Название'),
                Ссылка_на_материал=request.POST.get('Ссылка'),
                Курс_id=request.POST.get('Курс'),
                Тип_id=request.POST.get('Тип'),
            )
            return redirect('material_list')
    return render(request, 'main/MaterialList.html', {
        'материалы': Материалы_курсов.objects.all(),
        'курсы': Курсы.objects.all(),
        'типы': Типы_материалов.objects.all(),
    })
# --- Организации ---
def организации_view(request):
    if request.method == 'POST':
        if 'submit_организация' in request.POST:
            if Организации.objects.filter(ИНН=request.POST.get('ИНН')).exists():
                messages.warning(request, f"Организация с ИНН {request.POST.get('ИНН')} уже есть в базе данных!")
            elif Организации.objects.filter(ОГРН=request.POST.get('ОГРН')).exists():
                messages.warning(request, f"Организация с ОГРН {request.POST.get('ОГРН')} уже есть в базе данных!")
            else:
                Организации.objects.create(
                    Название=request.POST.get('Название'),
                    ИНН=request.POST.get('ИНН'),
                    ОГРН=request.POST.get('ОГРН'),
                    Телефон=request.POST.get('Телефон'),
                    Email=request.POST.get('Email'),
                )
                return redirect('organisation_list')
        elif 'submit_привязка' in request.POST:
            if Человек_организация.objects.filter(
                    Слушатель_id=request.POST.get('Слушатель'),
                    Организация_id=request.POST.get('Организация'),
                    Должность=request.POST.get('Должность'),
            ).exists():
                messages.warning(request, "Эта привязка уже есть в базе данных!")
            else:
                Человек_организация.objects.create(
                    Слушатель_id=request.POST.get('Слушатель'),
                    Организация_id=request.POST.get('Организация'),
                    Должность=request.POST.get('Должность'),
                )
                return redirect('organisation_list')
    return render(request, 'main/OrganisationList.html', {
        'слушатели': Слушатели.objects.all(),
        'организации': Организации.objects.all(),
        'привязки': Человек_организация.objects.all(),
    })

def delete_listener(request, id):
    listener = get_object_or_404(Слушатели, id=id)
    if Человек_группа.objects.filter(Слушатель=listener).exists() \
       or Человек_организация.objects.filter(Слушатель=listener).exists():
        messages.error(request, "Удаление отменено — есть зависимые записи.")
    else:
        listener.delete()
    return redirect('student_list')

def delete_group(request, id):
    group = get_object_or_404(Группы, id=id)
    if Человек_группа.objects.filter(Группа=group).exists():
        messages.error(request, "Удаление отменено — есть зависимые записи!")
    else:
        group.delete()
    return redirect('group_list')
def delete_group_linking(request, id):
    get_object_or_404(Человек_группа, id=id).delete()
    return redirect('group_list')
def delete_course(request, id):
    course = get_object_or_404(Курсы, id=id)
    if Группы.objects.filter(Курс=course).exists() \
       or Материалы_курсов.objects.filter(Курс=course).exists():
        messages.error(request, "Удаление отменено — есть зависимые записи!")
    else:
        course.delete()
    return redirect('course_list')

def delete_country(request, id):
<<<<<<< HEAD
    country = get_object_or_404(Страны, id=id)
    if Слушатели.objects.filter(Гражданство=country).exists():
        messages.error(request, "Удаление отменено — есть зависимые записи!")
    else:
        country.delete()
=======
    страна = get_object_or_404(Страны, id=id)
    if Слушатели.objects.filter(Гражданство=страна).exists() or Регионы.objects.filter(Страна=страна).exists():
        messages.error(request, "Удаление отменено — есть зависимые записи!")
    else:
        страна.delete()
>>>>>>> 2889caa (07.07.2025)
    return redirect('country_list')

def delete_material(request, id):
    get_object_or_404(Материалы_курсов, id=id).delete()
    return redirect('material_list')

def delete_organisation(request, id):
    org = get_object_or_404(Организации, id=id)
    if Человек_организация.objects.filter(Организация=org).exists():
        messages.error(request, "Удаление отменено — есть зависимые записи!")
    else:
        org.delete()
    return redirect('organisation_list')

def delete_org_linking(request, id):
    get_object_or_404(Человек_организация, id=id).delete()
    return redirect('organisation_list')

def today_date():
    month = Месяцы.objects.get(id=date.today().month).Название
    return f"{date.today().day} {month} {date.today().year} г."

def export_listeners_XLSX(request):
    try:
        # Создание Excel-файла
        wb = Workbook()
        ws = wb.active
        # Заголовки
        ws.append([
            "Фамилия", "Имя", "Отчество", "Дата рождения", "Пол", "Гражданство",
            "Серия паспорта", "Номер паспорта", "ИНН", "СНИЛС", "Телефон", "Email"
        ])
        # Данные
        for слушатель in Слушатели.objects.all():
            ws.append([
                слушатель.Фамилия,
                слушатель.Имя,
                слушатель.Отчество or "",
                слушатель.Дата_рождения.strftime("%d.%m.%Y") if слушатель.Дата_рождения else "",
                слушатель.Пол.Название if слушатель.Пол else "",
                слушатель.Гражданство.Краткое_название if слушатель.Гражданство else "",
                слушатель.Серия_паспорта or "",
                слушатель.Номер_паспорта or "",
                слушатель.ИНН or "",
                слушатель.Номер_СНИЛС or "",
                слушатель.Телефон or "",
                слушатель.Email or ""
            ])
            # Автоширина колонок
            for col in ws.columns:
                max_length = max(len(str(cell.value)) if cell.value else 0 for cell in col)
                ws.column_dimensions[get_column_letter(col[0].column)].width = max_length + 2
        # Отправка ответа
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = 'attachment;'

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        response.write(output.read())
        return response
    except Exception as e:
        return JsonResponse({"error": f"Ошибка при генерации Excel: {str(e)}"}, status=500)  # Уберите, если у вас работает CSRF через заголовки
def export_courses_XLSX(request):
    try:
        # Создание Excel-файла
        wb = Workbook()
        ws = wb.active
        # Заголовки
        ws.append(["Название курса", "Тип", "Количество часов"])
        # Данные
        for курс in Курсы.objects.all():
            ws.append([
                курс.Название,
                курс.Тип.Название,
                курс.Объём_часов
            ])
            # Автоширина колонок
            for col in ws.columns:
                max_length = max(len(str(cell.value)) if cell.value else 0 for cell in col)
                ws.column_dimensions[get_column_letter(col[0].column)].width = max_length + 2
        # Отправка ответа
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = 'attachment;'

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        response.write(output.read())
        return response
    except Exception as e:
        return JsonResponse({"error": f"Ошибка при генерации таблицы Excel: {str(e)}"}, status=500)
def export_groups_XLSX(request):
    try:
        # Создание Excel-файла
        wb = Workbook()
        ws = wb.active
        # Заголовки
        ws.append( [ "Слушатель", "Статус", "Группа", "Курс", "Тип курса"])
        # Данные
        for чел_группа in Человек_группа.objects.all():
            ws.append([
                f"{чел_группа.Слушатель.Фамилия} {чел_группа.Слушатель.Имя} {чел_группа.Слушатель.Отчество or ''} (ИНН: {чел_группа.Слушатель.ИНН})",
                чел_группа.Статус.Название,
                чел_группа.Группа_id,
                чел_группа.Группа.Курс.Название,
                чел_группа.Группа.Курс.Тип.Название,
            ])
        # Автоширина колонок (после заполнения)
        for column_cells in ws.columns:
            length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
            column_letter = get_column_letter(column_cells[0].column)
            ws.column_dimensions[column_letter].width = length + 2
        # Ответ пользователю
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = 'attachment;'

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        response.write(output.read())
        return response
    except Exception as e:
        return JsonResponse({"error": f"Ошибка при генерации  таблицы Excel: {str(e)}"}, status=500)

def export_organisations_XLSX(request):
    try:
        # Создание Excel-файла
        wb = Workbook()
        ws = wb.active
        # Заголовки
        ws.append([
            "Сотрудник", "Должность", "Название организации", "ИНН организации",
            "ОГРН организации", "Телефон организации", "Email организации"
        ])
        # Данные
        for чел_орг in Человек_организация.objects.all():
            ws.append([
                f"{чел_орг.Слушатель.Фамилия} {чел_орг.Слушатель.Имя} {чел_орг.Слушатель.Отчество or ''} (ИНН: {чел_орг.Слушатель.ИНН})",
                чел_орг.Должность,
                чел_орг.Организация.Название,
                чел_орг.Организация.ИНН,
                чел_орг.Организация.ОГРН,
                чел_орг.Организация.Телефон or '',
                чел_орг.Организация.Email or '',
            ])

        # Автоширина колонок (после заполнения)
        for column_cells in ws.columns:
            length = max(len(str(cell.value)) if cell.value else 0 for cell in column_cells)
            column_letter = get_column_letter(column_cells[0].column)
            ws.column_dimensions[column_letter].width = length + 2

        # Ответ пользователю
        response = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response["Content-Disposition"] = 'attachment;'
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        response.write(output.read())
        return response
    except Exception as e:
        return JsonResponse({"error": f"Ошибка при генерации таблицы Excel: {str(e)}"}, status=500)

def import_listeners_XLSX(request):
    listeners_count = Слушатели.objects.all().count()
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        try:
            wb = openpyxl.load_workbook(excel_file)
            sheet = wb.active

            for row in sheet.iter_rows(min_row=2, values_only=True):
                фамилия, имя, отчество, дата_рождения, пол, гражданство, серия, номер_паспорта, инн, снилс, телефон, email = row[:12]
                пол = Пол.objects.filter(Название__iexact=пол).first()
                гражданство = Страны.objects.filter(Краткое_название__iexact=гражданство).first()
                try:
                    # Преобразуем строку даты в объект date
                    if isinstance(дата_рождения, str):
                        дата_рождения = datetime.strptime(дата_рождения, '%d.%m.%Y').date()
                except Exception:
                    messages.error(request, f"Ошибка формата даты рождения: {дата_рождения}")
                    continue
                # Проверка: существует ли такой слушатель по уникальному набору полей (ИНН или паспорт)
                if Слушатели.objects.filter(
                    Фамилия=фамилия,
                    Имя=имя,
                    Отчество=отчество,
                    Дата_рождения=дата_рождения,
                    ИНН=инн
                ).exists():
                    continue
                if фамилия and имя and дата_рождения and пол and гражданство and номер_паспорта and инн:
                    Слушатели.objects.create(
                        Фамилия=фамилия,
                        Имя=имя,
                        Отчество=отчество,
                        Дата_рождения=дата_рождения,
                        Пол=пол,
                        Гражданство=гражданство,
                        Серия_паспорта=серия or '',
                        Номер_паспорта=номер_паспорта,
                        ИНН=инн,
                        Номер_СНИЛС=снилс or '',
                        Телефон=телефон or '',
                        Email=email or ''
                    )
            if listeners_count == Слушатели.objects.all().count():
                messages.warning(request, "Из таблицы Excel не было добавлено новых данных.")
            else:
                messages.success(request, f"Успешно добавлены новые данные.")
        except Exception as e:
            messages.error(request, f"Ошибка при открытии таблицы Excel: {e}")
    return redirect('student_list')
@csrf_exempt

def import_groups_XLSX(request):
    groups_count = Группы.objects.all().count()
    group_links_count = Человек_группа.objects.all().count()
    if request.method == "POST" and request.FILES.get("excel_file"):
        try:
            file = request.FILES["excel_file"]
            wb = openpyxl.load_workbook(file)
            ws = wb.active

            rows = list(ws.iter_rows(min_row=2, values_only=True))  # Пропускаем заголовок
            for row in rows:
                (
                    слушатель,
                    статус,
                    группа_id,
                    курс,
                    тип_курса,
                ) = row
                if "(ИНН:" not in слушатель:
                    continue  # пропускаем строку с неправильным форматом

                фио, инн = слушатель.split("(ИНН:")
                фио_parts = фио.strip().split()
                фамилия = фио_parts[0]
                имя = фио_parts[1] if len(фио_parts) > 1 else ""
                отчество = фио_parts[2] if len(фио_parts) > 2 else ""

                инн = инн.strip(") ").strip()
                курс = Курсы.objects.filter(Название=курс).first()
                статус = Статусы.objects.filter(Название=статус).first()
                try:
                    слушатель, _ = Слушатели.objects.get_or_create(
                        ИНН=инн,
                        defaults={
                            "Фамилия": фамилия,
                            "Имя": имя,
                            "Отчество": отчество,
                        },
                    )
                except Exception:
                    continue
                # --- Получение или создание группы ---
                группа, _ = Группы.objects.get_or_create(
                    id=группа_id,
                    defaults={
                        "Курс": курс,
                    },
                )
                # --- Добавляем привязку ---
                Человек_группа.objects.get_or_create(
                    Слушатель=слушатель,
                    Группа=группа,
                    Статус=статус
                )
            if groups_count == Группы.objects.all().count() and group_links_count == Человек_группа.objects.all().count():
                messages.warning(request, "Из таблицы Excel не было добавлено новых данных.")
            else:
                messages.success(request, "Успешно добавлены новые данные.")
            return redirect("group_list")  # редирект обратно на страницу
        except Exception as e:
            return JsonResponse(
                {"error": f"Ошибка при импорте из Excel: {str(e)}"}, status=400
            )
    return JsonResponse({"error": "Файл не был загружен."}, status=400)

def import_courses_XLSX(request):
    courses_count = Курсы.objects.all().count()
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        imported_count = 0
        try:
            wb = openpyxl.load_workbook(excel_file)
            sheet = wb.active
            for row in sheet.iter_rows(min_row=2, values_only=True):
                название, тип, число_часов = row[:3]
                тип = Типы_курсов.objects.filter(Название__iexact=тип).first()
                # Проверка: существует ли такой слушатель по уникальному набору полей (ИНН или паспорт)
                if Курсы.objects.filter( Название=название, Тип=тип, ).exists(): continue
                if название and тип and число_часов:
                    Курсы.objects.create(
                        Название=название,
                        Тип=тип,
                        Объём_часов=число_часов
                    )
            if courses_count == Курсы.objects.all().count():
                messages.warning(request, "Из таблицы Excel не было добавлено новых данных.")
            else:
                messages.success(request, f"Успешно добавлены новые данные.")
        except Exception as e:
            messages.error(request, f"Ошибка при открытии таблицы Excel: {e}")
    return redirect('course_list')

def import_organisations_XLSX(request):
    organisations_count = Группы.objects.all().count()
    organisation_links_count = Человек_группа.objects.all().count()
    if request.method == "POST" and request.FILES.get("excel_file"):
        try:
            file = request.FILES["excel_file"]
            wb = openpyxl.load_workbook(file)
            imported_count = 0
            ws = wb.active

            rows = list(ws.iter_rows(min_row=2, values_only=True))  # Пропускаем заголовок
            for row in rows:
                (
                    сотрудник_str,
                    должность,
                    название_организации,
                    инн_организации,
                    огрн,
                    телефон,
                    email,
                ) = row

                # --- Парсинг сотрудника ---
                # Формат: "Иванов Иван Иванович (ИНН: 1234567890)"
                if "(ИНН:" not in сотрудник_str:
                    continue  # пропускаем строку с неправильным форматом

                фио, инн = сотрудник_str.split("(ИНН:")
                фио_parts = фио.strip().split()
                фамилия = фио_parts[0]
                имя = фио_parts[1] if len(фио_parts) > 1 else ""
                отчество = фио_parts[2] if len(фио_parts) > 2 else ""

                инн = инн.strip(") ").strip()

                try:
                    слушатель, _ = Слушатели.objects.get_or_create(
                        ИНН=инн,
                        defaults={
                            "Фамилия": фамилия,
                            "Имя": имя,
                            "Отчество": отчество,
                        },
                    )
                except Exception:
                    continue
                # --- Получение или создание организации ---
                организация, _ = Организации.objects.get_or_create(
                    ИНН=инн_организации,
                    defaults={
                        "Название": название_организации,
                        "ОГРН": огрн,
                        "Телефон": телефон,
                        "Email": email,
                    },
                )
                # --- Добавляем привязку ---
                Человек_организация.objects.get_or_create(
                    Слушатель=слушатель,
                    Организация=организация,
                    defaults={"Должность": должность}
                )
            if organisations_count == Организации.objects.all().count() and organisation_links_count == Человек_организация.objects.all().count():
                messages.warning(request, "Из таблицы Excel не было добавлено новых данных.")
            else:
                messages.success(request, f"Успешно добавлены новые данные.")
            return redirect("organisation_list")  # редирект обратно на страницу
        except Exception as e:
            return JsonResponse(
                {"error": f"Ошибка при импорте из Excel: {str(e)}"}, status=400
            )
    return JsonResponse({"error": "Файл не был загружен."}, status=400)

def training_record(request):
    try:
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)

        # Заголовок организации
        org_paragraph = doc.add_paragraph()
        org_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = org_paragraph.add_run('Автономная некоммерческая организация дополнительного профессионального образования\n')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run = org_paragraph.add_run('«Сибирская академия профессионального обучения»')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True

        # Заголовок документа
        heading = doc.add_paragraph("Ведомость электронного обучения от " + today_date())
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].font.name = 'Times New Roman'
        heading.runs[0].font.size = Pt(14)

        # Получение ID групп из тела запроса
        body_unicode = request.body.decode('utf-8')
        ids = json.loads(body_unicode).get('ids', [])
        группы = Группы.objects.filter(id__in=ids).select_related('Курс', 'Курс__Тип')

        for группа in группы:
            # Название группы
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Группа №{группа.id} — {группа.Курс.Название} ({группа.Курс.Тип.Название})")
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            привязки = Человек_группа.objects.filter(Группа=группа).select_related('Слушатель', 'Статус')
            if not привязки.exists():
                paragraph = doc.add_paragraph("Слушателей нет")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.runs[0].font.name = 'Times New Roman'
                paragraph.runs[0].font.size = Pt(12)
                continue

            headers = ['Слушатель', 'Статус']
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'

            for i, text in enumerate(headers):
                cell = table.rows[0].cells[i]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            for привязка in привязки:
                слушатель = привязка.Слушатель
                row = table.add_row().cells
                values = [
                    f"{слушатель.Фамилия} {слушатель.Имя} {слушатель.Отчество or ''} (ИНН: {слушатель.ИНН})",
                    привязка.Статус.Название
                ]
                for i, text in enumerate(values):
                    cell = row[i]
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

        director = doc.add_paragraph("\nДиректор" + (' ' * 50) + "В.Н. Котлов")
        director.alignment = WD_ALIGN_PARAGRAPH.CENTER
        director.runs[0].font.name = 'Times New Roman'
        director.runs[0].font.size = Pt(14)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(buf.read(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment;'
        return response

    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=500)

def order(request):
    try:
        data = json.loads(request.body)
        ids = data.get("ids", [])
        param = data.get("type")  # 1 - зачисление, 2 - отчисление

        if not ids or len(ids) != 1:
            return JsonResponse({"error": "Выберите одну группу!"}, status=400)

        группа = Группы.objects.get(id=ids[0])
        participants = Человек_группа.objects.filter(Группа=группа)

        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_height = Mm(297)
        section.page_width = Mm(210)

        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)

        # Заголовок
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("Автономная некоммерческая организация\nдополнительного профессионального образования\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run = header.add_run("«Сибирская академия профессионального обучения»")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        participant_ids = sorted(str(p.Слушатель.id) for p in participants)
        hash_input = str(группа.id) + ''.join(participant_ids) + str(param) + today_date()
        order_hash = int(int(hashlib.md5(hash_input.encode('utf-8')).hexdigest(), 16)) % (2 ** 31)

        order, created = Приказы.objects.get_or_create(Hash=order_hash)
        order_number = f"{order.id}з" if param == 1 else f"{order.id}о"
        title = doc.add_paragraph(f"Приказ №{order_number}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.name = 'Times New Roman'
        title.runs[0].font.size = Pt(14)
        title.runs[0].bold = True

        date = doc.add_paragraph("г. Абакан" + ' ' * 50 + today_date())
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date.runs[0].font.name = 'Times New Roman'
        date.runs[0].font.size = Pt(12)

        if param == 1:
            intro_text = "О зачислении в число слушателей\nобразовательных программ\n"
            law_text = ( "В соответствии со ст. 53 Федерального закона «Об образовании в Российской Федерации» от 29.12.2012г. №273-ФЗ")
        else:
            intro_text = "Об отчислении слушателей\nи выдаче документов\nустановленного образца\n"
            law_text = ("В связи с успешным завершением обучения, в соответствии со ст. 61 Федерального закона «Об образовании в Российской Федерации» от 29.12.2012г. №273-ФЗ")
        intro = doc.add_paragraph(intro_text)
        intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
        intro.runs[0].font.name = 'Times New Roman'
        intro.runs[0].font.size = Pt(12)

        law = doc.add_paragraph(law_text)
        law.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        law.paragraph_format.first_line_indent = Mm(10)
        law.runs[0].font.name = 'Times New Roman'
        law.runs[0].font.size = Pt(12)

        pzv = doc.add_paragraph("ПРИКАЗЫВАЮ")
        pzv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pzv.runs[0].font.name = 'Times New Roman'
        pzv.runs[0].font.size = Pt(12)
        pzv.runs[0].bold = True

        if param == 1:
            body_text = f"1. Зачислить в число обучающихся по программе «{группа.Курс.Название}» следующих слушателей:"
        else:
            body_text = f"1. Слушателей, прошедших обучение по программе «{группа.Курс.Название}», отчислить и выдать документы установленного образца:"

        body = doc.add_paragraph(body_text)
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body.paragraph_format.first_line_indent = Mm(10)
        body.paragraph_format.space_after = Pt(4)
        body.runs[0].font.name = 'Times New Roman'
        body.runs[0].font.size = Pt(12)

        for p in participants:
            p.Статус_id = param
            p.save()
            p_line = doc.add_paragraph(f"{p.Слушатель.Фамилия} {p.Слушатель.Имя} {p.Слушатель.Отчество or ''}")
            p_line.paragraph_format.left_indent = Mm(10)
            p_line.paragraph_format.space_after = Pt(0)
            p_line.paragraph_format.space_before = Pt(0)
            p_line.runs[0].font.name = 'Times New Roman'
            p_line.runs[0].font.size = Pt(12)

        control = doc.add_paragraph("2. Контроль за исполнением настоящего приказа оставляю за собой.")
        control.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        control.paragraph_format.left_indent = Mm(10)
        control.paragraph_format.space_before = Pt(4)
        control.runs[0].font.name = 'Times New Roman'
        control.runs[0].font.size = Pt(12)

        director = doc.add_paragraph("\nДиректор" + (' ' * 50) + "В.Н. Котлов")
        director.alignment = WD_ALIGN_PARAGRAPH.CENTER
        director.runs[0].font.name = 'Times New Roman'
        director.runs[0].font.size = Pt(12)
        # Ответ клиенту
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        response = HttpResponse(
            buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment;"'
        return response

    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)}, status=500)

def generate_protocol_doc(организация):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    org_paragraph = doc.add_paragraph()
    org_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_paragraph.add_run('Автономная некоммерческая организация\nдополнительного профессионального образования\n')

    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = org_paragraph.add_run('«Сибирская академия профессионального обучения»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

    сотрудники = Человек_организация.objects.filter(Организация=организация)
    participant_ids = sorted(str(p.Слушатель.id) for p in сотрудники)
    hash_input = str(организация.id) + ''.join(participant_ids) + today_date()
    protocol_hash = int(int(hashlib.md5(hash_input.encode('utf-8')).hexdigest(), 16)) % (2 ** 31)

    protocol, created = Протоколы.objects.get_or_create(Hash=protocol_hash)
    heading = doc.add_paragraph(f"Протокол №{protocol.id} от {today_date()}\n\nзаседания комиссии по проверке знаний требований охраны труда работников\n\n{организация.Название}")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.size = Pt(12)

    decision = doc.add_paragraph("В соответствии с приказом директора АНО ДПО «САПО» №3 от 05 июля 2024 г.\nкомиссия в составе:\n\nПредседатель: Котлов В.Н. – Директор\nЧлены комиссии: Райков Т.Г. – Заместитель директора по развитию\nАзиханова Р.И. – Методист\n\nпровела проверку знаний требований охраны труда по программам:")
    decision.alignment = WD_ALIGN_PARAGRAPH.LEFT
    decision.paragraph_format.left_indent  = Mm(10)
    decision.runs[0].font.name = 'Times New Roman'
    decision.runs[0].font.size = Pt(12)

    headers = ['Ф.И.О.', 'Должность', 'Подпись проверяемого']
    for курс in Курсы.objects.all():
        чел_группа = Человек_группа.objects.filter(Группа__in=Группы.objects.filter(Курс=курс))
        слушатели = чел_группа.values_list('Слушатель__id', flat=True)
        сотрудники = Человек_организация.objects.filter(Слушатель__id__in=слушатели, Организация=организация)

        if сотрудники.count() == 0: continue

        course = doc.add_paragraph(f"{курс.id}. {курс.Название}, в объёме {курс.Объём_часов} часов.")
        course.paragraph_format.space_before = Pt(10)
        course.paragraph_format.space_after = Pt(0)
        course.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        course.paragraph_format.first_line_indent = Mm(10)
        course.runs[0].font.name = 'Times New Roman'
        course.runs[0].font.size = Pt(12)

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, text in enumerate(headers):
            paragraph = table.rows[0].cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        for i, привязка in enumerate(сотрудники):
            row = table.add_row().cells
            values = [
                f"{привязка.Слушатель.Фамилия} {привязка.Слушатель.Имя} {привязка.Слушатель.Отчество or ''}",
                привязка.Должность, ''
            ]
            for j, text in enumerate(values):
                cell = row[j]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    director = doc.add_paragraph("\nПредседатель" + (' ' * 55) + "В.Н. Котлов")
    director.alignment = WD_ALIGN_PARAGRAPH.LEFT
    director.paragraph_format.left_indent = Mm(10)
    director.runs[0].font.name = 'Times New Roman'
    director.runs[0].font.size = Pt(12)

    members = doc.add_paragraph("Члены комиссии" + (' ' * 50) + "Т.Г. Райков, Р.И. Азиханова")
    members.alignment = WD_ALIGN_PARAGRAPH.LEFT
    members.paragraph_format.left_indent = Mm(10)
    members.runs[0].font.name = 'Times New Roman'
    members.runs[0].font.size = Pt(12)
    return doc
def generate_certificate_doc(организация):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(15)
    section.right_margin = Mm(30)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    привязки = Человек_организация.objects.filter(Организация=организация).select_related('Слушатель')
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    for привязка in привязки:
        row = table.add_row().cells
        # Левая ячейка
        left = row[0]
        left_para = left.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = left_para.add_run(
            "Автономная некоммерческая организация\n"
            "дополнительного профессионального образования\n"
            "«Сибирская академия профессионального обучения»\n"
        )
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True

        run = left_para.add_run("(Лицензия Министерства образования и науки РХ №Л035-01237-19/00257225 от 18.05.2017 г.)\n\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(6)

        run = left_para.add_run(f"УДОСТОВЕРЕНИЕ № {привязка.id}\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        # ФИО таблица
        fio_table = left.add_table(rows=3, cols=2)
        fio_table.style = 'Table Grid'

        fields = [("Фамилия", привязка.Слушатель.Фамилия),
                  ("Имя", привязка.Слушатель.Имя),
                  ("Отчество", привязка.Слушатель.Отчество or '')]

        for i, (label, value) in enumerate(fields):
            label_cell = fio_table.cell(i, 0)
            value_cell = fio_table.cell(i, 1)

            label_para = label_cell.paragraphs[0]
            value_para = value_cell.paragraphs[0]

            run = label_para.add_run(label)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            run = value_para.add_run(value)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True

        # Должность и организация
        p = left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{привязка.Должность}\n\n{организация.Название}\n")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

        # Правая ячейка
        right = row[1]
        title = right.paragraphs[0]
        title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = title.add_run("Прошел обучение и проверку знаний требований охраны труда по программам:")
        title.paragraph_format.space_after = Pt(4)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

        # Список курсов
        курс_группы = Человек_группа.objects.filter(Слушатель=привязка.Слушатель).select_related("Группа__Курс")
        курсы_уникальные = set()
        for i, запись in enumerate(курс_группы):
            запись.Статус_id = 2
            if запись.Группа.Курс.id in курсы_уникальные:
                continue
            курсы_уникальные.add(запись.Группа.Курс.id)
            string = f"{i + 1}. {запись.Группа.Курс.Название}, в объёме {запись.Группа.Курс.Объём_часов} часов"
            if i < len(курс_группы) - 1: string += ';'
            else: string += '.'
            p = right.add_paragraph(string)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Mm(5)
            run = p.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)

        # Протокол
        protocol = right.add_paragraph()
        protocol.alignment = WD_ALIGN_PARAGRAPH.LEFT
        protocol.paragraph_format.space_before = Pt(6)
        protocol.paragraph_format.space_after = Pt(6)
        run1 = protocol.add_run("Протокол №")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(10)
        run2 = protocol.add_run(f"{Протоколы.objects.last().id} от {today_date()}")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)
        run2.bold = True
        # Подпись директора
        director = right.add_paragraph("Директор" + (' ' * 50) + "Котлов В.Н.")
        director.alignment = WD_ALIGN_PARAGRAPH.LEFT
        director.runs[0].font.name = 'Times New Roman'
        director.runs[0].font.size = Pt(10)
    return doc


def generate_vedomost_doc(организация):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    org_paragraph = doc.add_paragraph()
    org_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = org_paragraph.add_run('Автономная некоммерческая организация\nдополнительного профессионального образования\n')

    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run = org_paragraph.add_run('«Сибирская академия профессионального обучения»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    heading = doc.add_paragraph("Ведомость выдачи документов от " + today_date())
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.name = 'Times New Roman'
    heading.runs[0].font.size = Pt(14)

    # Заголовок группы
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Организация: {организация.Название} (ИНН: {организация.ИНН}, ОГРН: {организация.ОГРН})")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    привязки = Человек_организация.objects.filter(Организация=организация).select_related('Слушатель')

    headers = ['ФИО слушателя', 'Должность', 'Номер бланка удостоверения (свидетельства)', 'Номер протокола',
               'Дата выдачи', 'Подпись']

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    for i, text in enumerate(headers):
        paragraph = table.rows[0].cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    протокол = str(Протоколы.objects.last().id) if Протоколы.objects.exists() else '0'

    for привязка in привязки:
        row = table.add_row().cells
        values = [
            f"{привязка.Слушатель.Фамилия} {привязка.Слушатель.Имя} {привязка.Слушатель.Отчество or ''}",
            привязка.Должность, str(привязка.id), протокол, today_date(), ''
        ]
        for i, text in enumerate(values):
            cell = row[i]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    director = doc.add_paragraph("\nДиректор" + (' ' * 50) + "В.Н. Котлов")
    director.alignment = WD_ALIGN_PARAGRAPH.CENTER
    director.runs[0].font.name = 'Times New Roman'
    director.runs[0].font.size = Pt(14)
    return doc
def generate_certificates_zip(request):
    if request.method == "POST":
        body = json.loads(request.body)
        ids = body.get("ids", [])
        if not ids:
            return JsonResponse({"error": "Нет выбранных ID."}, status=400)

        организация = Организации.objects.filter(id__in=ids).first()
        сотрудники = Человек_организация.objects.filter(Организация=организация)
        if not сотрудники.exists():
            return JsonResponse({"error": "Нет сотрудников!"}, status=400)

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Протокол
            doc2 = generate_protocol_doc(организация)
            doc_bytes2 = BytesIO()
            doc2.save(doc_bytes2)
            zip_file.writestr(f"Протокол №{Протоколы.objects.last().id}.docx", doc_bytes2.getvalue())
            # Удостоверения
            doc1 = generate_certificate_doc(организация)
            doc_bytes1 = BytesIO()
            doc1.save(doc_bytes1)
            zip_file.writestr("Удостоверения.docx", doc_bytes1.getvalue())
            # Ведомость выдачи сертификатов
            doc3 = generate_vedomost_doc(организация)
            doc_bytes3 = BytesIO()
            doc3.save(doc_bytes3)
            zip_file.writestr(f"Ведомость выдачи сертификатов.docx", doc_bytes3.getvalue())

        zip_buffer.seek(0)
        response = HttpResponse(zip_buffer.read(), content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename=Документы_по_удостоверениям.zip'
        return response
<<<<<<< HEAD
    return JsonResponse({"error": "Метод не поддерживается."}, status=405)
=======
    return JsonResponse({"error": "Метод не поддерживается."}, status=405)

def get_regions(request):
    страна_id = request.GET.get('Страна')
    регионы = Регионы.objects.filter(Страна_id=страна_id).order_by('Название')
    return JsonResponse(list(регионы.values('id', 'Название')), safe=False)

def get_districts(request):
    region_id = request.GET.get('region_id')
    районы = Районы.objects.filter(Регион_id=region_id).order_by('Название')
    return JsonResponse(list(районы.values('id', 'Название')), safe=False)

def get_places(request):
    region_id = request.GET.get('region_id')
    district_id = request.GET.get('district_id')
    if district_id:
        places = Населенные_пункты.objects.filter(Район_id=district_id)
    else:
        places = Населенные_пункты.objects.filter(Регион_id=region_id)
    places = places.order_by('Название')
    return JsonResponse(list(places.values('id', 'Название')), safe=False)

def get_postcodes(request):
    place_id = request.GET.get('place_id')
    postcodes = Почтовые_индексы.objects.filter(Населенный_пункт_id=place_id).order_by('Индекс')
    return JsonResponse(list(postcodes.values('id', 'Индекс')), safe=False)

def get_streets(request):
    place_id = request.GET.get('place_id')
    улицы = Улицы.objects.filter(Населенный_пункт_id=place_id).select_related('Название').order_by('Название__Название')
    result = [{'id': улица.id, 'Название': улица.Название.Название} for улица in улицы]
    return JsonResponse(result, safe=False)
def адреса_слушателей_view(request):
    if request.method == "POST":
        try:
            слушатель = get_object_or_404(Слушатели, id=request.POST.get('Слушатель'))
            улица = get_object_or_404(Улицы, id=request.POST.get('Улица'))
            индекс = get_object_or_404(Почтовые_индексы, id=request.POST.get('Почтовый_индекс'))

            адрес, created = Адреса.objects.get_or_create(
                улица=улица,
                Почтовый_индекс=индекс,
                Номер_дома=request.POST.get('Номер_дома')
            )
            if Адреса_слушателей.objects.filter(Слушатель=слушатель).exists():
                messages.warning(request, "Этот человек уже есть в таблице адресов!")
            else:
                Адреса_слушателей.objects.create(Слушатель=слушатель, Адрес=адрес)
            return redirect('listener_addresses')
        except IntegrityError:
            messages.error(request, "Для этого человека уже задан адрес!")
        except Exception as e:
            messages.error(request, f"Ошибка при добавлении адреса: {str(e)}")
            return redirect('listener_addresses')
    адреса = Адреса_слушателей.objects.select_related(
        'Слушатель', 'Адрес__улица__Название',
        'Адрес__улица__Населенный_пункт__Район__Регион__Страна',
        'Адрес__Почтовый_индекс'
    )
    return render(request, 'main/ListenerAddresses.html', {
        'страны': Страны.objects.all(),
        'адреса': адреса,
        'слушатели': Слушатели.objects.all()
    })
def адреса_организаций_view(request):
    if request.method == "POST":
        try:
            организация = get_object_or_404(Организации, id=request.POST.get('Организация'))
            улица = get_object_or_404(Улицы, id=request.POST.get('Улица'))
            индекс = get_object_or_404(Почтовые_индексы, id=request.POST.get('Почтовый_индекс'))

            адрес, created = Адреса.objects.get_or_create(
                улица=улица,
                Почтовый_индекс=индекс,
                Номер_дома=request.POST.get('Номер_дома')
            )
            if Адреса_организаций.objects.filter(Организация=организация).exists():
                messages.warning(request, "Этот человек уже есть в таблице адресов!")
            else:
                Адреса_организаций.objects.create(Организация=организация, Адрес=адрес)
            return redirect('organisation_addresses')
        except IntegrityError:
            messages.error(request, "Для этой организации уже задан адрес!")
        except Exception as e:
            messages.error(request, f"Ошибка при добавлении адреса: {str(e)}")
            return redirect('organisation_addresses')
    адреса = Адреса_организаций.objects.select_related(
        'Организация', 'Адрес__улица__Название',
        'Адрес__улица__Населенный_пункт__Район__Регион__Страна',
        'Адрес__Почтовый_индекс'
    )
    return render(request, 'main/OrganisationAddresses.html', {
        'страны': Страны.objects.all(),
        'адреса': адреса,
        'организации': Организации.objects.all()
    })
def delete_listener_address(request, id):
    obj = get_object_or_404(Адреса_слушателей, id=id)
    obj.delete()
    return redirect('listener_addresses')

def delete_organisation_address(request, id):
    get_object_or_404(Адреса_организаций, id=id).delete()
    return redirect('organisation_addresses')
>>>>>>> 2889caa (07.07.2025)
